import asyncio
from pathlib import Path
from threading import Lock
from typing import ClassVar
from uuid import UUID

from docx import Document
from fastapi import HTTPException, status
from loguru import logger
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics, ttfonts
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import Settings
from app.models.enums import ExportFormat, ExportStatus
from app.models.export import TranscriptExport
from app.repositories.export_repository import ExportRepository
from app.repositories.transcript_repository import TranscriptRepository
from app.schemas.auth import CurrentUser
from app.schemas.export import ExportJobResponse, ExportRead
from app.services.storage_service import StorageService


class ExportService:
    _pdf_font_lock: ClassVar[Lock] = Lock()
    _pdf_fonts_registered: ClassVar[bool] = False
    _pdf_fonts: ClassVar[dict[str, str]] = {
        "latin": "TranscribeAI-NotoSans",
        "arabic": "TranscribeAI-NotoSansArabic",
        "devanagari": "TranscribeAI-NotoSansDevanagari",
        "cjk": "TranscribeAI-NotoSansJP",
    }

    def __init__(self, session: AsyncSession, settings: Settings) -> None:
        self.settings = settings
        self.storage = StorageService(settings)
        self.export_repository = ExportRepository(session)
        self.transcript_repository = TranscriptRepository(session)

    async def list_exports(self, user: CurrentUser) -> list[ExportRead]:
        exports = await self.export_repository.list_for_user(user.id)
        logger.info("[EXPORT] Listed - user={} count={}", user.id, len(exports))
        return [ExportRead.model_validate(item) for item in exports]

    async def enqueue_export(self, user: CurrentUser, transcript_id: UUID, format: ExportFormat) -> ExportJobResponse:
        transcript = await self.transcript_repository.get_for_user(transcript_id, user.id)
        if not transcript:
            logger.warning("[EXPORT] Transcript not found - user={} transcript_id={} format={}", user.id, transcript_id, format.value)
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Transcript not found")

        export = await self.export_repository.create(
            TranscriptExport(
                transcript_id=transcript.id,
                user_id=user.id,
                format=format.value,
                status=ExportStatus.QUEUED.value,
            )
        )
        logger.info("[EXPORT] Record created - user={} export_id={} format={}", user.id, export.id, format.value)

        job_id = str(export.id)
        if self.settings.background_tasks_enabled:
            from app.workers.export_tasks import generate_export_task

            logger.info("[EXPORT] Submitting to queue - export_id={} format={}", export.id, export.format)
            async_result = generate_export_task.delay(str(export.id))
            job_id = async_result.id
            logger.info("[EXPORT] Queued - export_id={} job_id={}", export.id, job_id)
        else:
            logger.info("[EXPORT] Generating inline (background tasks disabled) - export_id={}", export.id)
            await self.generate_export_file(export.id)

        return ExportJobResponse(job_id=job_id, export_id=export.id, status=export.status, format=export.format)

    async def generate_export_file(self, export_id: UUID) -> TranscriptExport:
        export = await self.export_repository.get_by_id(export_id)
        if not export:
            logger.error("[EXPORT] Export not found - export_id={}", export_id)
            raise ValueError(f"Export not found: {export_id}")
        transcript = await self.transcript_repository.get_by_id(export.transcript_id)
        if not transcript:
            logger.error("[EXPORT] Transcript not found - export_id={} transcript_id={}", export_id, export.transcript_id)
            raise ValueError(f"Transcript not found: {export.transcript_id}")

        directory = self.storage.transcript_export_directory(transcript.id)
        path = directory / f"transcript.{export.format}"
        logger.info("[EXPORT] Generating file - export_id={} format={} path={}", export.id, export.format, path)
        await asyncio.to_thread(self._write_export_file, path, export.format, transcript.title, transcript.text)
        size_bytes = path.stat().st_size
        updated = await self.export_repository.update_status(
            export, status=ExportStatus.COMPLETED.value, storage_path=str(path), size_bytes=size_bytes
        )
        logger.info("[EXPORT] File ready - export_id={} format={} size={}B", export.id, export.format, size_bytes)
        return updated

    async def get_or_create_export_path(self, user: CurrentUser, transcript_id: UUID, format: ExportFormat) -> Path:
        transcript = await self.transcript_repository.get_for_user(transcript_id, user.id)
        if not transcript:
            logger.warning("[EXPORT] Transcript not found for download - user={} transcript_id={} format={}", user.id, transcript_id, format.value)
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Transcript not found")

        directory = self.storage.transcript_export_directory(transcript.id)
        path = directory / f"transcript.{format.value}"
        if not path.exists():
            logger.info("[EXPORT] File missing, regenerating - transcript_id={} format={}", transcript_id, format.value)
            await asyncio.to_thread(self._write_export_file, path, format.value, transcript.title, transcript.text)
        logger.info("[EXPORT] Download ready - user={} transcript_id={} format={}", user.id, transcript_id, format.value)
        return path

    def _write_export_file(self, path: Path, format: str, title: str, text: str) -> None:
        if format == ExportFormat.TXT.value:
            path.write_text(text, encoding="utf-8")
            return
        if format == ExportFormat.PDF.value:
            self._write_pdf(path, title, text)
            return
        if format == ExportFormat.DOCX.value:
            self._write_docx(path, title, text)
            return
        raise ValueError(f"Unsupported export format: {format}")

    def _write_pdf(self, path: Path, title: str, text: str) -> None:
        self._register_pdf_fonts()
        document = canvas.Canvas(str(path), pagesize=letter)
        width, height = letter
        left_margin = 72
        right_margin = width - 72
        max_width = width - 144
        y = height - 72

        title_font = self._select_pdf_font(title)
        self._draw_pdf_line(
            document,
            title[:90],
            font_name=title_font,
            font_size=14,
            left=left_margin,
            right=right_margin,
            y=y,
        )
        y -= 30

        for paragraph in text.splitlines() or [text]:
            body_font = self._select_pdf_font(paragraph)
            for line in self._wrap_pdf_text(paragraph, body_font, 10, max_width):
                if y < 72:
                    document.showPage()
                    y = height - 72
                self._draw_pdf_line(
                    document,
                    line,
                    font_name=body_font,
                    font_size=10,
                    left=left_margin,
                    right=right_margin,
                    y=y,
                )
                y -= 14
        document.save()

    def _write_docx(self, path: Path, title: str, text: str) -> None:
        document = Document()
        document.add_heading(title, level=1)
        for paragraph in text.splitlines() or [text]:
            document.add_paragraph(paragraph)
        document.save(path)

    @classmethod
    def _register_pdf_fonts(cls) -> None:
        if cls._pdf_fonts_registered:
            return

        with cls._pdf_font_lock:
            if cls._pdf_fonts_registered:
                return

            font_directory = Path(__file__).resolve().parents[1] / "assets" / "fonts"
            font_files = {
                cls._pdf_fonts["latin"]: font_directory / "NotoSans.ttf",
                cls._pdf_fonts["arabic"]: font_directory / "NotoSansArabic.ttf",
                cls._pdf_fonts["devanagari"]: font_directory / "NotoSansDevanagari.ttf",
                cls._pdf_fonts["cjk"]: font_directory / "NotoSansJP.ttf",
            }
            missing = [str(path) for path in font_files.values() if not path.is_file()]
            if missing:
                raise RuntimeError(f"PDF font assets are missing: {', '.join(missing)}")

            for font_name, font_path in font_files.items():
                pdfmetrics.registerFont(TTFont(font_name, str(font_path), shapable=True))

            cls._pdf_fonts_registered = True
            logger.info("[EXPORT] Unicode PDF fonts registered - directory={}", font_directory)

    @classmethod
    def _select_pdf_font(cls, value: str) -> str:
        script_counts = {
            "arabic": sum(
                "\u0600" <= character <= "\u06ff"
                or "\u0750" <= character <= "\u077f"
                or "\u08a0" <= character <= "\u08ff"
                or "\ufb50" <= character <= "\ufdff"
                or "\ufe70" <= character <= "\ufeff"
                for character in value
            ),
            "devanagari": sum(
                "\u0900" <= character <= "\u097f"
                or "\ua8e0" <= character <= "\ua8ff"
                for character in value
            ),
            "cjk": sum(
                "\u3040" <= character <= "\u30ff"
                or "\u3400" <= character <= "\u4dbf"
                or "\u4e00" <= character <= "\u9fff"
                or "\uac00" <= character <= "\ud7af"
                for character in value
            ),
        }
        script = max(script_counts, key=script_counts.get)
        return cls._pdf_fonts[script] if script_counts[script] else cls._pdf_fonts["latin"]

    def _wrap_pdf_text(
        self,
        value: str,
        font_name: str,
        font_size: int,
        max_width: float,
    ) -> list[str]:
        if not value:
            return [""]

        lines: list[str] = []
        current = ""
        for token in value.split(" "):
            candidate = token if not current else f"{current} {token}"
            if current and pdfmetrics.stringWidth(candidate, font_name, font_size) > max_width:
                lines.extend(self._split_wide_pdf_token(current, font_name, font_size, max_width))
                current = token
            else:
                current = candidate

        if current:
            lines.extend(self._split_wide_pdf_token(current, font_name, font_size, max_width))
        return lines or [""]

    def _split_wide_pdf_token(
        self,
        value: str,
        font_name: str,
        font_size: int,
        max_width: float,
    ) -> list[str]:
        if pdfmetrics.stringWidth(value, font_name, font_size) <= max_width:
            return [value]

        chunks: list[str] = []
        current = ""
        for character in value:
            candidate = f"{current}{character}"
            if current and pdfmetrics.stringWidth(candidate, font_name, font_size) > max_width:
                chunks.append(current)
                current = character
            else:
                current = candidate
        if current:
            chunks.append(current)
        return chunks

    def _draw_pdf_line(
        self,
        document: canvas.Canvas,
        value: str,
        *,
        font_name: str,
        font_size: int,
        left: float,
        right: float,
        y: float,
    ) -> None:
        document.setFont(font_name, font_size)
        shaped_value = ttfonts.shapeStr(value, font_name, font_size, force=True)
        if self._contains_rtl_text(value):
            document.drawRightString(right, y, shaped_value)
        else:
            document.drawString(left, y, shaped_value)

    def _contains_rtl_text(self, value: str) -> bool:
        return any(
            "\u0590" <= character <= "\u08ff"
            or "\ufb1d" <= character <= "\ufdff"
            or "\ufe70" <= character <= "\ufeff"
            for character in value
        )
